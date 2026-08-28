#!/usr/bin/env python3
"""Generate diploma documentation DOCX for Student Management System."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / 'docs' / 'screenshots'

OUT_PATHS = [
    ROOT / 'docs' / 'SMS_Diploma_Documentation_Shpak_Maksym.docx',
    Path('/Users/maksym/Diploma/Documentation for the diploma project - Student Management System - Shpak Maksym.docx'),
]


def set_normal_style(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def add_figure(doc: Document, filename: str, caption: str) -> None:
    """Embed a PNG from docs/screenshots, or fall back to the caption alone."""
    path = SHOTS / filename
    if path.exists():
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(str(path), width=Cm(15.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    run = cap.add_run(caption)
    run.italic = True
    _set_run_font(run, size=Pt(10))


def _set_run_font(run, *, bold: bool = False, size: Pt = Pt(12)) -> None:
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = size
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def _set_cell_width(cell, width_dxa: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        from docx.oxml import OxmlElement
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')


def add_basic_info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """2-column table like BloomTime sample (label | content)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.autofit = False

    left_w, right_w = 2800, 6526  # dxa — same proportions as BloomTime

    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        _set_cell_width(left, left_w)
        _set_cell_width(right, right_w)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        left.text = ''
        lp = left.paragraphs[0]
        lr = lp.add_run(label)
        _set_run_font(lr, bold=True)

        right.text = ''
        first = True
        for line in value.split('\n'):
            rp = right.paragraphs[0] if first else right.add_paragraph()
            first = False
            # Bold Backend:/Frontend:/Infrastructure: labels inside cell
            if line.rstrip().endswith(':') and line.rstrip() in {
                'Backend:', 'Frontend:', 'Infrastructure:'
            }:
                rr = rp.add_run(line)
                _set_run_font(rr, bold=True)
            else:
                rr = rp.add_run(line)
                _set_run_font(rr, bold=False)

    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    set_normal_style(doc)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # Title page
    for line in [
        'School of Computer Science & Technologies',
        'Field of study: Computer Science',
        '',
        'Maksym Shpak',
        '',
        'Design and Implementation of a Web-Based Student Management System '
        'for Educational Institutions',
        '',
        'Documentation for the diploma project',
        'prepared under the direction of',
        'Marcin Kacprowicz',
        '',
        '',
        'Warsaw, 2026 r.',
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14 if 'Design and Implementation' in line else 12)
            if 'Design and Implementation' in line or line == 'Maksym Shpak':
                run.bold = True

    doc.add_page_break()

    # 1. Basic information — table (same layout as BloomTime sample)
    add_heading(doc, '1. Basic information', 1)

    tech_list = (
        'Backend:\n'
        'Python 3.12, Django 6, Django Auth (session-based), Gunicorn, WhiteNoise, '
        'dj-database-url, PostgreSQL (production) / SQLite (development)\n'
        'Frontend:\n'
        'Django Templates, HTML5, CSS custom properties, vanilla JavaScript, '
        'Chart.js, Turbo, Bootstrap Icons, dashboard JSON endpoint (/api/dashboard/)\n'
        'Infrastructure:\n'
        'Render + Neon (build.sh: migrate + collectstatic + seed_demo), '
        'English/Polish i18n, HTTPS-oriented security settings when DEBUG=False, '
        'live at https://thesms.me (Render backup: '
        'https://diploma-project-sms-management-system.onrender.com)'
    )

    add_basic_info_table(doc, [
        (
            'Project name',
            'Design and Implementation of a Web-Based Student Management System '
            'for Educational Institutions — Student Management System (SMS)',
        ),
        (
            'Purpose of the project',
            'The project addresses a recurring problem in small and medium educational '
            'units: academic records for students, courses, grades and attendance are still '
            'often kept in spreadsheets or fragmented paper files. This leads to duplicated '
            'data, weak access control, and slow reporting for teachers and administrators. '
            'SMS centralises these processes in a single web application with role-based '
            'access (Administrator, Teacher), structured CRUD workflows, dashboard '
            'analytics, CSV export. Staff accounts are created by an administrator '
            '(there is no public registration). Students are records in the system, '
            'not user accounts. The system is implemented '
            'as a complete engineering deliverable — not only a conceptual design.',
        ),
        (
            'Brief description of the project',
            'Student Management System is a full-stack Django web application that stores '
            'and manages academic data related to students, courses, enrolments, grades and '
            'attendance. Administrators maintain the institutional catalogue of students and '
            'courses; teachers manage grades and attendance for their own courses; students '
            'are academic records, not login accounts. The presentation layer uses Django templates '
            'with custom CSS and JavaScript; the application layer implements business rules '
            'and RBAC; the data layer uses a relational schema (SQLite locally, PostgreSQL in '
            'production via DATABASE_URL). The application is live at '
            'https://thesms.me (Render backup: '
            'https://diploma-project-sms-management-system.onrender.com).',
        ),
        (
            'Competitor analysis',
            'Microsoft Excel / Google Sheets: easy to start, but no relational integrity, '
            'no role-based permissions, and high risk of accidental edits. '
            'Moodle / Canvas: powerful LMS platforms with many modules, but heavy to deploy '
            'and operate for a small faculty office that only needs student-course-grade-'
            'attendance administration. '
            'Commercial university ERP suites: broad feature sets and high cost, often '
            'oversized for training centres or single departments. '
            'SMS differentiator: a lightweight, open academic engineering project focused on '
            'core administrative workflows, transparent Django architecture suitable for '
            'defence and further extension, zero proprietary licence fees, and explicit '
            'teacher-centred dashboards with JSON analytics API.',
        ),
        (
            'List of technologies used',
            tech_list,
        ),
        (
            'Description of the technology stack and justification of selected technologies',
            'The stack follows a classic three-tier architecture: Browser (HTML/CSS/JS) → '
            'Django views & forms (business logic, authentication, RBAC) → Relational database. '
            'Django was selected because it provides batteries-included authentication, CSRF '
            'protection, ORM migrations, and a dedicated Teachers screen for issuing lecturer '
            'accounts — essential for an academic records system delivered within an engineering '
            'diploma timeframe. Session authentication '
            'fits server-rendered pages better than JWT (no separate SPA). PostgreSQL is preferred '
            'in production for ACID transactions and concurrent multi-user access; SQLite keeps '
            'local development simple. WhiteNoise serves compressed static assets behind Gunicorn, '
            'which matches typical PaaS deployments without a separate Nginx container. '
            'There is no public registration: an administrator creates teacher accounts '
            'and issues credentials. The Teachers screen inside SMS is used to set a '
            'forgotten password, which matches a faculty-office workflow better than '
            'self-service email recovery.',
        ),
    ])

    # 2. Key issues
    add_heading(doc, '2. Key issues related to the implementation of the project', 1)

    add_heading(doc, '2.1 Role-Based Access Control (RBAC)', 2)
    add_para(
        doc,
        'The custom User model extends AbstractUser with a role field '
        '(admin / teacher). View decorators enforce permissions: administrators '
        'manage all records; teachers see only students, grades and attendance related '
        'to their own courses. A guessed URL outside that scope returns 404. '
        'Students are not users — they have no login. This prevents privilege '
        'escalation while keeping URLs simple.',
    )
    add_para(doc, 'Code snippet (decorator, simplified):', bold=True)
    add_para(
        doc,
        "def role_required(allowed_roles):\n"
        "    def decorator(view_func):\n"
        "        def wrapper(request, *args, **kwargs):\n"
        "            if request.user.role not in allowed_roles:\n"
        "                return redirect('dashboard')\n"
        "            return view_func(request, *args, **kwargs)\n"
        "        return wrapper\n"
        "    return decorator",
    )

    add_heading(doc, '2.2 Relational Academic Schema', 2)
    add_para(
        doc,
        'The schema comprises six core models: User, Student, Course, Enrollment, Grade, '
        'Attendance. Enrolment links students and courses (many-to-many with status). A '
        'Grade is one weighted component of an enrolment (coursework, midterm, final exam '
        'or retake); Enrollment.final_grade is their weighted mean on the 2.0–5.0 scale '
        'with letter mapping, and assigned_by records which teacher last saved the mark. '
        'Attendance is unique per (enrolment, date). Foreign keys and unique constraints '
        'protect referential integrity at the database level.',
    )

    add_heading(doc, '2.3 Dashboard Analytics and REST JSON Endpoint', 2)
    add_para(
        doc,
        'The HTML dashboard aggregates KPIs (students, courses, enrolments, average grade) '
        'and chart series (grade distribution, enrolments per course). Endpoint '
        'GET /api/dashboard/ returns the same metrics as JSON, scoped by role, so charts '
        'or future SPA clients can refresh data without reloading the page.',
    )
    add_para(doc, 'Example response fields:', bold=True)
    add_para(
        doc,
        "{'role': 'teacher', 'total_students': …, 'grade_distribution': {'A': …}, "
        "'course_labels': [...], 'course_data': [...]}",
    )

    add_heading(doc, '2.4 Staff Account Provisioning', 2)
    add_para(
        doc,
        'The application has no public registration. An administrator creates teacher '
        'accounts (username, role, initial password) on the Teachers screen inside SMS '
        'and issues those credentials out of band. Lecturers sign in with the issued '
        'username and password; students are data records, not users. If a teacher '
        'forgets a password, the administrator sets a new one on that same screen. '
        'This matches the real workflow of a faculty office and avoids an email-based '
        'recovery flow that would be unused without self-service sign-up.',
    )

    add_heading(doc, '2.5 Production Configuration and Security Hardening', 2)
    add_para(
        doc,
        'Configuration is driven by environment variables: SECRET_KEY, DEBUG, ALLOWED_HOSTS, '
        'DATABASE_URL, DJANGO_ADMIN_PASSWORD. When DEBUG=False the application enforces HTTPS '
        'redirects, secure cookies, HSTS, and refuses insecure default SECRET_KEY values. '
        'DJANGO_ADMIN_PASSWORD replaces the published demo password on the admin account at '
        'deploy time, so that password is not left on the public internet. WhiteNoise '
        'serves hashed/compressed static files. This allows the same codebase to run locally '
        'on SQLite and on a hosted PostgreSQL instance without code changes. '
        'The public service is https://thesms.me (Render backup: '
        'https://diploma-project-sms-management-system.onrender.com). '
        'A live demonstration uses the teacher account prof.martinez / demo1234; a second '
        'teacher, prof.chen / demo1234, shows that each lecturer sees only their own courses.',
    )

    add_heading(doc, '2.6 Internationalisation (English / Polish)', 2)
    add_para(
        doc,
        'The interface is bilingual. Templates, model choices, flash messages and CSV '
        'headers go through Django i18n. The language switch keeps the current page and '
        'stores the choice in the session. The compiled Polish catalogue (django.mo) is '
        'committed so the deploy image does not depend on gettext. LocalizationTests '
        'guard that English and Polish pages actually differ.',
    )

    # 3. Screenshots
    add_heading(doc, '3. Screenshots, visualizations, etc.', 1)
    add_para(
        doc,
        'The figures below are captures of the running Student Management System '
        '(local interface, English locale, light theme unless noted).',
    )
    add_figure(
        doc, 'login.png',
        'Fig. 1 — Login page. Accounts are issued by an administrator; there is no public registration.',
    )
    add_figure(
        doc, 'dashboard-light.png',
        'Fig. 2 — Administrator dashboard: KPI tiles and Chart.js analytics.',
    )
    add_figure(
        doc, 'students-list.png',
        'Fig. 3 — Student register with search, programme and status filters.',
    )
    add_figure(
        doc, 'courses-list.png',
        'Fig. 4 — Course catalogue: code, semester, teacher assignment and capacity.',
    )
    add_figure(
        doc, 'grades-list.png',
        'Fig. 5 — Grade register: weighted components, letter bands and the teacher who saved the mark.',
    )
    add_figure(
        doc, 'attendance-list.png',
        'Fig. 6 — Attendance register with present / absent / late statuses.',
    )
    add_figure(
        doc, 'teachers-list.png',
        'Fig. 7 — Teachers screen: the administrator issues lecturer accounts and can set a new password.',
    )
    add_figure(
        doc, 'student-detail.png',
        'Fig. 8 — Student profile: enrolments, grades and attendance in one view.',
    )

    # 4. Conclusions
    add_heading(doc, '4. Conclusions and development prospects', 1)
    add_para(
        doc,
        'The primary goal — to design and implement a functional web-based Student '
        'Management System for educational administration — was achieved. The application '
        'covers authentication with RBAC, student and course administration, enrolments, '
        'weighted grade components, bulk attendance marking, a bilingual EN/PL interface, '
        'analytics dashboard, CSV export, and production-oriented configuration.',
    )
    add_para(doc, 'All core objectives were met:', bold=True)
    for item in [
        'Role-based access for Administrator and Teacher (students are records, not users)',
        'Relational schema with integrity constraints for academic entities',
        'CRUD workflows with form validation and flash messages',
        'Weighted grade components, authorship audit, bulk class attendance',
        'Dashboard KPIs and role-scoped REST endpoint /api/dashboard/',
        'Bilingual interface (English / Polish, 308 translatable strings)',
        'Staff accounts provisioned by an administrator (no public registration)',
        'Automated test suite (80 tests) covering models, auth, API and CRUD',
        'Deployment-ready settings (Gunicorn, WhiteNoise, Postgres via DATABASE_URL)',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_para(
        doc,
        'The project demonstrates that a production-capable academic administration system '
        'can be delivered by a single developer within an engineering diploma timeframe, '
        'using a coherent Django monolith suitable for oral examination and further maintenance.',
    )

    add_para(doc, 'Future development priorities:', bold=True)
    for item in [
        'Student portal — a third role with read-only access to own results',
        'PDF report generation for transcripts and attendance sheets',
        'Email notifications on new grades and absences',
        'Documented REST API with token authentication for mobile clients',
        'Two-factor login (TOTP) for administrator accounts',
        'Timetabling with room and slot conflict detection',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # 5. Bibliography
    add_heading(doc, '5. Bibliography / Resources', 1)

    add_para(doc, 'Official Documentation', bold=True)
    refs = [
        '[1] Django Software Foundation — Django Documentation — https://docs.djangoproject.com',
        '[2] Django Software Foundation — Django Authentication System — https://docs.djangoproject.com/en/stable/topics/auth/',
        '[3] Python Software Foundation — Python 3 Documentation — https://docs.python.org/3/',
        '[4] The PostgreSQL Global Development Group — PostgreSQL Documentation — https://www.postgresql.org/docs/',
        '[5] SQLite Consortium — SQLite Documentation — https://www.sqlite.org/docs.html',
        '[6] Whitenoise — WhiteNoise Documentation — https://whitenoise.readthedocs.io',
        '[7] Gunicorn — Gunicorn Documentation — https://docs.gunicorn.org',
        '[8] Render — Deploy Django — https://render.com/docs/deploy-django',
        '[9] MDN Web Docs — HTML / CSS / HTTP — https://developer.mozilla.org',
        '[10] OWASP Foundation — OWASP Top Ten — https://owasp.org/www-project-top-ten/',
        '[11] Jazzband — dj-database-url — https://github.com/jazzband/dj-database-url',
        '[12] Bootstrap Icons — https://icons.getbootstrap.com',
        '[13] Uniwersytet VIZJA / AEH — study programme materials for Computer Science (engineering profile)',
    ]
    for r in refs:
        add_para(doc, r)

    return doc


def main() -> None:
    doc = build()
    for path in OUT_PATHS:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        print('Wrote', path)


if __name__ == '__main__':
    main()
